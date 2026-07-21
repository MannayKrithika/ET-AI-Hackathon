"""
Text extraction — Step 3.

One function per file type, all returning the same shape, so the upload
route doesn't need to know or care which file type it's dealing with:

    {
        "success": bool,
        "text": str,
        "pages_processed": int | None,   # PDF only, else None
        "char_count": int,
        "error": str | None,
    }

Supported now: PDF, DOCX, XLSX, CSV.
Still skipped: PNG, JPG, scanned PDFs — that's OCR, which comes later.
"""

from pathlib import Path

import fitz  # PyMuPDF
import docx  # python-docx
import pandas as pd

from text_cleaning import clean_text_encoding
from tabular_utils import load_sheet_with_detected_header


def _result(success, text="", pages_processed=None, error=None, page_word_counts=None):
    return {
        "success": success,
        "text": text,
        "pages_processed": pages_processed,
        "char_count": len(text),
        "error": error,
        # Word count of each page, in order — lets chunking map a chunk's
        # position in the full text back to the page it came from.
        "page_word_counts": page_word_counts,
    }


def extract_pdf_text(file_path: Path) -> dict:
    """Open PDF -> read each page -> join all text -> return."""
    try:
        doc = fitz.open(str(file_path))
        page_texts = [clean_text_encoding(page.get_text()) for page in doc]
        pages_processed = doc.page_count
        doc.close()

        full_text = "\n\n".join(page_texts).strip()
        # Word count per page, computed the same way (str.split()) as the
        # chunk word counts downstream, so the two line up.
        page_word_counts = [len(p.split()) for p in page_texts]

        return _result(
            True,
            full_text,
            pages_processed=pages_processed,
            page_word_counts=page_word_counts,
        )
    except Exception as e:
        return _result(False, error=str(e))


def extract_docx_text(file_path: Path) -> dict:
    """Open Word doc -> read every paragraph -> join -> return."""
    try:
        document = docx.Document(str(file_path))
        paragraphs = [clean_text_encoding(p.text) for p in document.paragraphs if p.text.strip()]

        # Word docs often carry tables too — pull those in as plain rows.
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(clean_text_encoding(cell.text) for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs).strip()
        return _result(True, full_text)
    except Exception as e:
        return _result(False, error=str(e))


def _rows_to_records(df, label: str) -> str:
    """
    Convert a DataFrame into boxed "Field: Value" record blocks — much
    easier for an LLM to parse later than one long space-joined row.

        ========== Breakdown Record ==========

        Ticket_ID: BRK-8000
        Asset_ID: PMP-KS-29
        ...

        ======================================
    """
    df = df.fillna("")
    border = "=" * 38
    records = []
    for _, row in df.iterrows():
        lines = [f"{col}: {row[col]}" for col in df.columns]
        block = f"========== {label} Record ==========\n\n" + "\n".join(lines) + f"\n\n{border}"
        records.append(block)

    return "\n\n".join(records)


def _find_header_row(raw_df, max_scan: int = 10) -> int:
    """
    Some spreadsheets have a merged title row (and sometimes a blank row)
    sitting above the real header row, e.g.:

        Row 0: "Equipment Breakdown History & RCA"   <- title, mostly empty cells
        Row 1: (blank)
        Row 2: "Ticket ID", "Asset ID", "Equipment Name", ...   <- real header

        Naively treating row 0 as the header produces "Unnamed: 1",
        "Unnamed: 2", etc. for every real column.

    Heuristic: the real header row and the data rows below it both use
    (roughly) the full column width. A title row, by contrast, only has
    one or two non-null cells (the rest are NaN from merged cells). So we
    take the most common "non-null cell count" across all rows as the
    sheet's real width, then scan top-down for the first row that matches
    that width — that's the header.
    """
    non_null_counts = raw_df.notna().sum(axis=1)

    if len(non_null_counts) == 0:
        return 0

    full_width = non_null_counts.mode().iloc[0]

    scan_limit = min(max_scan, len(raw_df))
    for i in range(scan_limit):
        if non_null_counts.iloc[i] >= full_width:
            return i

    return 0  # fallback: nothing matched, assume row 0 is the header


def extract_xlsx_text(file_path: Path) -> dict:
    """
    Open workbook -> for each sheet, detect the real header row (skipping
    any title/blank rows above it) -> convert each data row to a labeled
    record.
    """
    try:
        raw_sheets = pd.read_excel(file_path, sheet_name=None, header=None, dtype=str)
        blocks = []

        for sheet_name, raw_df in raw_sheets.items():
            if raw_df.empty:
                continue

            header_idx = _find_header_row(raw_df)
            headers = [
                str(h).strip().replace(" ", "_") if pd.notna(h) else f"Column_{i+1}"
                for i, h in enumerate(raw_df.iloc[header_idx])
            ]

            data_df = raw_df.iloc[header_idx + 1:].copy()
            data_df.columns = headers
            data_df = data_df.dropna(how="all")  # skip fully blank rows

            records_text = _rows_to_records(data_df, label=sheet_name)
            blocks.append(records_text)

        full_text = "\n\n".join(blocks).strip()
        return _result(True, full_text)
    except Exception as e:
        return _result(False, error=str(e))


def extract_csv_text(file_path: Path) -> dict:
    """Open CSV -> convert each row to a labeled record (same pattern as Excel)."""
    try:
        df = pd.read_csv(file_path, dtype=str)
        full_text = _rows_to_records(df, label=file_path.stem.replace("_", " ")).strip()
        return _result(True, full_text)
    except Exception as e:
        return _result(False, error=str(e))


# Dispatch table: extension -> extractor function.
# PNG/JPG aren't here yet — they're still uploaded and saved, just not
# extracted until OCR is added.
EXTRACTORS = {
    ".pdf": extract_pdf_text,
    ".docx": extract_docx_text,
    ".xlsx": extract_xlsx_text,
    ".csv": extract_csv_text,
}


def extract_text(file_path: Path) -> dict | None:
    """Returns None if this file type isn't supported for extraction yet."""
    extractor = EXTRACTORS.get(file_path.suffix.lower())
    if extractor is None:
        return None
    return extractor(file_path)
